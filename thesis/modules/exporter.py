"""Export thesis to multiple formats: DOCX, PDF, LaTeX, Markdown."""

import os
from pathlib import Path
from datetime import datetime


class ThesisExporter:
    """Export thesis chapters to various formats."""

    def __init__(self, db, config=None):
        self.db = db
        self.config = config or {}

    def export_docx(self, output_path: str, metadata: dict = None) -> str:
        """Export thesis to DOCX format."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        meta = metadata or {}
        doc = Document()

        # Page setup
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(4)
            section.right_margin = Cm(3)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Title page
        doc.add_paragraph()
        doc.add_paragraph()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(meta.get('title', 'Judul Skripsi'))
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'

        doc.add_paragraph()
        author = doc.add_paragraph()
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author.add_run(meta.get('author', 'Nama Mahasiswa'))
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

        if meta.get('nim'):
            nim = doc.add_paragraph()
            nim.alignment = WD_ALIGN_PARAGRAPH.CENTER
            nim.add_run(f"NIM: {meta['nim']}").font.size = Pt(12)

        if meta.get('university'):
            uni = doc.add_paragraph()
            uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
            uni.add_run(meta['university']).font.size = Pt(14)

        doc.add_paragraph()
        year = doc.add_paragraph()
        year.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year.add_run(str(meta.get('year', datetime.now().year))).font.size = Pt(12)

        doc.add_page_break()

        # Get chapters from database
        chapters = self.db.execute(
            "SELECT * FROM chapters ORDER BY chapter_number"
        ).fetchall()

        for ch in chapters:
            heading = doc.add_heading(ch['title'], level=1)
            for run in heading.runs:
                run.font.name = 'Times New Roman'

            content = ch['content'] or ''
            for para_text in content.split('\n\n'):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.first_line_indent = Cm(1.27)
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

            doc.add_page_break()

        # Bibliography
        refs = self.db.execute(
            "SELECT * FROM references_table ORDER BY citation_key"
        ).fetchall()

        if refs:
            doc.add_heading('Daftar Pustaka', level=1)
            for ref in refs:
                citation = ref.get('formatted_apa', ref.get('title', ''))
                p = doc.add_paragraph(citation)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.left_indent = Cm(1.27)
                p.paragraph_format.first_line_indent = Cm(-1.27)

        doc.save(output_path)
        return output_path

    def export_markdown(self, output_path: str, metadata: dict = None) -> str:
        """Export thesis to Markdown format."""
        meta = metadata or {}
        lines = []

        lines.append(f"# {meta.get('title', 'Judul Skripsi')}\n")
        lines.append(f"**{meta.get('author', 'Nama Mahasiswa')}**\n")
        if meta.get('nim'):
            lines.append(f"NIM: {meta['nim']}\n")
        lines.append(f"*{meta.get('university', 'Universitas')} — {meta.get('year', datetime.now().year)}*\n")
        lines.append("---\n")

        chapters = self.db.execute(
            "SELECT * FROM chapters ORDER BY chapter_number"
        ).fetchall()

        for ch in chapters:
            lines.append(f"\n## {ch['title']}\n")
            lines.append(ch.get('content', '') + "\n")

        refs = self.db.execute(
            "SELECT * FROM references_table ORDER BY citation_key"
        ).fetchall()

        if refs:
            lines.append("\n## Daftar Pustaka\n")
            for ref in refs:
                lines.append(f"- {ref.get('formatted_apa', ref.get('title', ''))}")

        Path(output_path).write_text('\n'.join(lines), encoding='utf-8')
        return output_path

    def export_latex(self, output_path: str, metadata: dict = None) -> str:
        """Export thesis to LaTeX format."""
        meta = metadata or {}
        template = r"""\documentclass[12pt,a4paper]{report}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{times}
\usepackage{setspace}
\usepackage[margin=2.54cm]{geometry}
\usepackage{graphicx}
\usepackage{hyperref}
\usepackage{natbib}

\onehalfspacing

\title{""" + meta.get('title', 'Judul Skripsi') + r"""}
\author{""" + meta.get('author', 'Nama Mahasiswa') + r"""}
\date{""" + str(meta.get('year', datetime.now().year)) + r"""}

\begin{document}

\maketitle
\tableofcontents
\newpage

"""
        chapters = self.db.execute(
            "SELECT * FROM chapters ORDER BY chapter_number"
        ).fetchall()

        for ch in chapters:
            template += f"\\chapter{{{ch['title']}}}\n"
            template += (ch.get('content', '') or '') + "\n\n"

        template += r"\bibliographystyle{apalike}" + "\n"
        template += r"\bibliography{references}" + "\n"
        template += r"\end{document}" + "\n"

        Path(output_path).write_text(template, encoding='utf-8')
        return output_path

    def export_bibtex(self, output_path: str) -> str:
        """Export all references as BibTeX."""
        refs = self.db.execute("SELECT * FROM references_table").fetchall()
        entries = []

        for ref in refs:
            entry = f"@article{{{ref.get('citation_key', 'unknown')},\n"
            if ref.get('title'):
                entry += f"  title = {{{ref['title']}}},\n"
            if ref.get('authors'):
                entry += f"  author = {{{ref['authors']}}},\n"
            if ref.get('year'):
                entry += f"  year = {{{ref['year']}}},\n"
            if ref.get('journal'):
                entry += f"  journal = {{{ref['journal']}}},\n"
            if ref.get('doi'):
                entry += f"  doi = {{{ref['doi']}}},\n"
            entry += "}\n"
            entries.append(entry)

        Path(output_path).write_text('\n'.join(entries), encoding='utf-8')
        return output_path
